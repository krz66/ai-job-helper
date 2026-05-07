from flask import Flask, render_template, request, jsonify
from ai_service import ai_agent
import sys
import os
from docx import Document
import threading
import time

if getattr(sys, 'frozen', False):
    template_folder = os.path.join(sys._MEIPASS, 'templates')
else:
    template_folder = 'templates'

app = Flask(__name__, template_folder=template_folder)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024

chat_history = []
active_requests = {}
cancel_flags = {}


def extract_text_from_file(file):
    filename = file.filename.lower()
    try:
        if filename.endswith('.docx'):
            doc = Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            return text.strip()
        elif filename.endswith('.txt'):
            return file.read().decode('utf-8', errors='ignore')
        elif filename.endswith('.md'):
            return file.read().decode('utf-8', errors='ignore')
        else:
            raise Exception("不支持的文件类型，请上传docx/txt/md格式")
    except Exception as e:
        raise Exception(f"文件读取失败：{str(e)}")


@app.route('/')
def index():
    return render_template('index.html', history=chat_history[-30:])


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "没有文件上传"})
    file = request.files['file']
    set_as_default = request.form.get('set_as_default', 'true') == 'true'
    if file.filename == '':
        return jsonify({"error": "未选择文件"})
    try:
        content = extract_text_from_file(file)
        if set_as_default and content.strip():
            ai_agent.set_default_resume(content)
        return jsonify({
            "content": content,
            "char_count": len(content)
        })
    except Exception as e:
        return jsonify({"error": str(e)})


@app.route('/set_default_resume', methods=['POST'])
def set_default_resume():
    data = request.json
    resume_text = data.get('resume', '')
    if not resume_text.strip():
        return jsonify({"error": "简历内容不能为空"})
    ai_agent.set_default_resume(resume_text)
    return jsonify({"status": "success", "message": f"默认简历已更新（{len(resume_text)}字符）"})


@app.route('/get_default_resume', methods=['GET'])
def get_default_resume():
    resume = ai_agent.get_default_resume()
    return jsonify({
        "has_default": bool(resume),
        "char_count": len(resume),
        "preview": resume[:200] + "..." if len(resume) > 200 else resume
    })


@app.route('/clear_default_resume', methods=['POST'])
def clear_default_resume():
    ai_agent.set_default_resume("")
    return jsonify({"status": "success", "message": "默认简历已清除"})


@app.route('/upload_and_process', methods=['POST'])
def upload_and_process():
    if 'file' not in request.files:
        return jsonify({"error": "没有文件上传"})
    file = request.files['file']
    mode = request.form.get('mode', 'extract')
    set_as_default = request.form.get('set_as_default', 'true') == 'true'
    if file.filename == '':
        return jsonify({"error": "未选择文件"})
    try:
        content = extract_text_from_file(file)
        if set_as_default and content.strip():
            ai_agent.set_default_resume(content)
        if mode == 'extract':
            result = ai_agent.extract_resume(content)
        elif mode == 'polish':
            result = ai_agent.polish_resume(content)
        else:
            result = "不支持的模式"
        chat_history.append({
            "mode": mode,
            "input": f"[上传文件: {file.filename}]",
            "output": result
        })
        return jsonify({"result": result})
    except Exception as e:
        return jsonify({"error": str(e)})


@app.route('/process', methods=['POST'])
def process():
    data = request.json
    mode = data.get('mode')
    text = data.get('text', '')
    text2 = data.get('text2', '')
    request_id = str(time.time())
    cancel_flags[request_id] = False

    def process_request():
        try:
            ai_agent.reset()
            if cancel_flags.get(request_id, False):
                active_requests[request_id] = "请求已取消"
                return

            if mode == 'extract':
                result = ai_agent.extract_resume(text)
            elif mode == 'polish':
                result = ai_agent.polish_resume(text)
            elif mode == 'match':
                if text2.strip():
                    result = ai_agent.match_job(text, text2)
                else:
                    default_resume = ai_agent.get_default_resume()
                    if not default_resume.strip():
                        result = "⚠️ 未检测到默认简历，请先上传简历再进行匹配！"
                    else:
                        result = ai_agent.match_job(text, default_resume)
            elif mode == 'interview':
                result = ai_agent.gen_interview(text)
            elif mode == 'qa':
                result = ai_agent.qa_ask(text)
            else:
                result = "未知模式"

            if cancel_flags.get(request_id, False):
                active_requests[request_id] = "请求已取消"
                return

            chat_history.append({
                "mode": mode,
                "input": text if not text2 else f"{text}\n---\n{text2}",
                "output": result
            })
            active_requests[request_id] = result
        except Exception as e:
            active_requests[request_id] = f"处理失败：{str(e)}"
        finally:
            if request_id in cancel_flags:
                del cancel_flags[request_id]

    thread = threading.Thread(target=process_request)
    thread.start()
    timeout = 180
    start = time.time()
    while request_id not in active_requests and time.time() - start < timeout:
        time.sleep(0.2)
    if request_id in active_requests:
        result = active_requests.pop(request_id)
    else:
        result = "⏰ 请求超时，请稍后重试"
        if request_id in cancel_flags:
            del cancel_flags[request_id]
    return jsonify({"result": result})


@app.route('/cancel', methods=['POST'])
def cancel():
    data = request.json
    request_id = data.get('request_id')
    if request_id in cancel_flags:
        cancel_flags[request_id] = True
    return jsonify({"status": "canceled"})


if __name__ == '__main__':
    import webbrowser

    port = 5000
    webbrowser.open(f'http://127.0.0.1:{port}')
    app.run(host='127.0.0.1', port=port, debug=False, threaded=True)